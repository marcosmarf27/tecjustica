"""Biblioteca TecJustica DOCX Builder — padrão institucional.

Gera relatórios processuais no padrão visual de bancas de investimento e
consultorias top-tier (Goldman, JP Morgan, McKinsey, BCG), adaptado para o
contexto judicial brasileiro. Usa python-docx para escrever o DOCX e
LibreOffice headless para conversão DOCX -> PDF.

Design system — Sobriedade Institucional
=========================================

Paleta oficial TecJustica · Edição Executiva:
    NAVY          #12223F   primária (deep navy jurídico)
    NAVY_DEEP     #0A1428   topo de capa, barras institucionais
    OCRE          #A67C2E   acento único (dourado envelhecido)
    OCRE_LIGHT    #C9A96E   linhas divisórias sutis, hairlines
    CREAM         #F6F1E6   fundo de capa, metadata grid
    CREAM_DARK    #ECE4D1   linhas alternadas sutis
    BODY          #2B2B2B   corpo (nunca preto puro)
    MUTED         #6B6B6B   labels, legendas, paginação
    HAIRLINE      #D4C9A8   hairlines dourados 0.25pt

Tipografia editorial:
    EB Garamond     — display + body (serifa transicional clássica)
    IBM Plex Sans   — labels, eyebrows, metadata em CAPS com tracking
    IBM Plex Mono   — document ID, datas, numerais, paginação

Regras tipográficas:
    - Labels em UPPERCASE com character spacing amplo (tracking +200)
    - Corpo sempre em serifa, nunca sans
    - Datas em formato longo: "14 de abril de 2026"
    - Numerais de página em mono: "01 / 48"
    - Quebras de linha manuais no título da capa (nunca deixar o Word quebrar)

Uso:
    from docx_builder import Report

    r = Report(
        titulo="Relatório de Análise Processual",
        subtitulo="Parecer técnico sobre cumprimento de sentença",
        classificacao="RESERVADO",
        numero_documento="TJ-CE / 2026 / 024",
        metadata=[
            ("AUTOS", "0001234-56.2024.8.06.0001"),
            ("CLASSE", "Cumprimento de sentença"),
            ("ÓRGÃO JULGADOR", "4ª Vara Cível · Fortaleza/CE"),
            ("RELATOR(A)", "Juiz(a) de Direito"),
            ("DATA", "14 de abril de 2026"),
        ],
        autor="TecJustica · Assessoria Judicial",
    )
    r.add_cover()
    r.add_toc()
    r.add_heading("Sumário executivo", level=1, numeral="01")
    r.add_paragraph("...")
    r.save("relatorio.docx")
    r.export_pdf("relatorio.docx", "relatorio.pdf")
"""

from __future__ import annotations

import subprocess
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as _Doc
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.table import _Cell


# =========================================================================
# Paleta institucional
# =========================================================================

NAVY = RGBColor(0x12, 0x22, 0x3F)
NAVY_DEEP = RGBColor(0x0A, 0x14, 0x28)
OCRE = RGBColor(0xA6, 0x7C, 0x2E)
OCRE_LIGHT = RGBColor(0xC9, 0xA9, 0x6E)
CREAM = RGBColor(0xF6, 0xF1, 0xE6)
BODY = RGBColor(0x2B, 0x2B, 0x2B)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CREAM_DARK = RGBColor(0xEC, 0xE4, 0xD1)

HEX_NAVY = "12223F"
HEX_NAVY_DEEP = "0A1428"
HEX_OCRE = "A67C2E"
HEX_OCRE_LIGHT = "C9A96E"
HEX_CREAM = "F6F1E6"
HEX_CREAM_DARK = "ECE4D1"
HEX_HAIRLINE = "D4C9A8"
HEX_WHITE = "FFFFFF"
HEX_ROW_ALT = "FBF8F0"
HEX_CODE_BG = "F4EFE2"

FONT_DISPLAY = "Caladea"
FONT_BODY = "EB Garamond 12"
FONT_SANS = "IBM Plex Sans"
FONT_MONO = "IBM Plex Mono"


# =========================================================================
# Helpers XML de baixo nível
# =========================================================================

def _set_cell_bg(cell: _Cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(
    cell: _Cell,
    *,
    top: tuple[str, int] | None = None,
    bottom: tuple[str, int] | None = None,
    left: tuple[str, int] | None = None,
    right: tuple[str, int] | None = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    sides = {"top": top, "left": left, "bottom": bottom, "right": right}
    for side, value in sides.items():
        b = OxmlElement(f"w:{side}")
        if value is None:
            b.set(qn("w:val"), "nil")
        else:
            color, sz = value
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _set_table_width(table, width_cm: float) -> None:
    """Define a largura total da tabela em centímetros (OOXML w:tblW)."""
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def _set_table_indent(table, left_cm: float = 0) -> None:
    """Remove a indentação default da tabela (OOXML w:tblInd)."""
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblInd"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(int(left_cm * 567)))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)


def _set_row_height(row, height_cm: float, rule: str = "exact") -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:trHeight"))
    if existing is not None:
        tr_pr.remove(existing)
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), str(int(height_cm * 567)))  # 567 twips per cm
    height.set(qn("w:hRule"), rule)
    tr_pr.append(height)


def _set_cell_margins(cell: _Cell, top: int = 100, bottom: int = 100, left: int = 140, right: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _add_left_border(paragraph, color: str, size: int = 24) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _tracking(run, value: int) -> None:
    """Character spacing em twentieths of a point (OOXML w:spacing).

    Valores razoáveis:
        0       texto normal (não mexer)
        20      labels pequenos em caps (1pt de extra)
        40      títulos display em caps (2pt)
        60      eyebrows bem respirados (3pt)
        80      nome rotacionado em barra lateral (4pt)

    Valores acima de 80 geram espaços absurdos entre letras.
    """
    r_pr = run._r.get_or_add_rPr()
    existing = r_pr.find(qn("w:spacing"))
    if existing is not None:
        r_pr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(value))
    r_pr.append(spacing)


def _set_run_font(run, name: str) -> None:
    run.font.name = name
    r_pr = run._r.get_or_add_rPr()
    rfonts = r_pr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        r_pr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _insert_field(run, instruction: str, placeholder: str = " ") -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder_el = OxmlElement("w:t")
    placeholder_el.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder_el)
    run._r.append(fld_end)


def _styled_run(
    paragraph,
    text: str,
    *,
    font: str = FONT_BODY,
    size: int = 11,
    color: RGBColor = BODY,
    bold: bool = False,
    italic: bool = False,
    tracking: int = 0,
    uppercase: bool = False,
):
    run = paragraph.add_run(text.upper() if uppercase else text)
    _set_run_font(run, font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if tracking:
        _tracking(run, tracking)
    return run


# =========================================================================
# Report
# =========================================================================

@dataclass
class Report:
    titulo: str
    subtitulo: str = ""
    classificacao: str = "DOCUMENTO RESERVADO"
    numero_documento: str = ""
    autor: str = "TecJustica · Assessoria Judicial"
    metadata: Sequence[tuple[str, str]] = field(default_factory=list)
    eyebrow: str = "RELATÓRIO PROCESSUAL"
    orgao: str = "TecJustica · Assessoria Judicial com IA"
    doc: _Doc = field(default_factory=Document)

    def __post_init__(self) -> None:
        self._configure_styles()
        self._configure_page()
        self._configure_header_footer()

    # ---- Setup ----

    def _configure_styles(self) -> None:
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = FONT_BODY
        normal.font.size = Pt(11)
        normal.font.color.rgb = BODY
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.35

    def _configure_page(self) -> None:
        # A seção inicial é usada PELA CAPA: A4, margens zeradas, sem header/footer
        # O conteúdo após a capa recebe uma nova seção com margens normais
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.header_distance = Cm(0)
        section.footer_distance = Cm(0)

    def _configure_header_footer(self) -> None:
        """Header/footer ficam vazios na seção da capa (sec 0).

        A seção de conteúdo (criada depois da capa via _add_content_section)
        é que recebe o header/footer institucional completo.
        """
        section = self.doc.sections[0]
        # Limpar parágrafos default do header/footer para não desenhar nada na capa
        section.header.paragraphs[0].clear()
        section.footer.paragraphs[0].clear()

    def _setup_content_section_header_footer(self, section) -> None:
        """Aplica header e footer institucional a uma seção de conteúdo."""
        # Desvincular do header/footer herdado da seção anterior (capa)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header = section.header
        p = header.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)
        _styled_run(p, "TECJUSTICA", font=FONT_SANS, size=8, color=NAVY, bold=True, tracking=40, uppercase=True)
        p.add_run("\t")
        _styled_run(p, self.classificacao, font=FONT_SANS, size=7, color=MUTED, tracking=30, uppercase=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(7.75), WD_ALIGN_PARAGRAPH.CENTER)
        tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)

        _styled_run(
            p, self.numero_documento or "TECJUSTICA · ASSESSORIA JUDICIAL",
            font=FONT_MONO, size=7, color=MUTED, tracking=20, uppercase=True,
        )
        p.add_run("\t")
        _styled_run(p, "TECJUSTICA · EDIÇÃO EXECUTIVA", font=FONT_SANS, size=7, color=MUTED, tracking=40, uppercase=True)
        p.add_run("\t")

        page_run = p.add_run()
        _set_run_font(page_run, FONT_MONO)
        page_run.font.size = Pt(7)
        page_run.font.color.rgb = NAVY
        page_run.bold = True
        _insert_field(page_run, "PAGE")
        _styled_run(p, " / ", font=FONT_MONO, size=7, color=MUTED)
        total_run = p.add_run()
        _set_run_font(total_run, FONT_MONO)
        total_run.font.size = Pt(7)
        total_run.font.color.rgb = NAVY
        total_run.bold = True
        _insert_field(total_run, "NUMPAGES")

    # ---- Cover (Editorial com barra lateral vertical) ----

    def add_cover(self) -> None:
        """Capa editorial estilo Weekly Report Template.

        Comportamento: a primeira seção do documento (section 0) foi
        configurada em __post_init__ com A4, margens zeradas e header/footer
        vazios — ideal para a capa sangrada. Depois da capa, criamos uma
        nova seção de conteúdo com margens normais e header/footer
        institucional via _add_content_section().

        Antigo comentário sobre barra lateral + eyebrow — mantido apenas
        para histórico:

        Layout:
            +-----+------------------------------------+
            |  B  |  (topo: respiro)                   |
            |  A  |  eyebrow · número do documento     |
            |  R  |  ─────────── (filete fino ocre)    |
            |  R  |                                    |
            |  A  |  TÍTULO                            |
            |     |  GIGANTESCO                        |
            |  V  |  (Caladea 72pt bold)               |
            |  E  |                                    |
            |  R  |  Subtítulo em itálico              |
            |  T  |                                    |
            |  I  |  ════════════ (filete ocre 2pt)    |
            |  C  |                                    |
            |  A  |  (metadata 3x2)                    |
            |  L  |  AUTOS      CLASSE       VALOR     |
            |     |  ÓRGÃO      RELATOR      DATA      |
            |  N  |                                    |
            |  A  |  ─────────── (filete fino)         |
            |  V  |  PREPARADO POR                     |
            |  Y  |  autor                             |
            +-----+------------------------------------+

        A barra lateral ocupa altura total da página com cor navy sólido
        e contém texto vertical rotacionado (TECJUSTICA · ASSESSORIA JUDICIAL).
        """
        # Limpar qualquer parágrafo vazio que python-docx criou no início
        body = self.doc.element.body
        for child in list(body):
            if child.tag == qn("w:p") and not (child.text or "").strip():
                has_runs = child.findall(qn("w:r"))
                if not has_runs:
                    body.remove(child)
                break

        self._build_cover_grid()

        # Criar nova seção para o conteúdo (margens normais + header/footer)
        from docx.enum.section import WD_SECTION
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.page_width = Cm(21.0)
        new_section.page_height = Cm(29.7)
        new_section.top_margin = Mm(28)
        new_section.bottom_margin = Mm(24)
        new_section.left_margin = Mm(26)
        new_section.right_margin = Mm(32)
        new_section.header_distance = Mm(15)
        new_section.footer_distance = Mm(15)
        self._setup_content_section_header_footer(new_section)

    def _build_cover_grid(self) -> None:
        """Capa editorial estilo Weekly Report Template.

        Layout:
            +----------------------------------------+
            |                                        |
            |  EYEBROW · DOC ID                      |
            |                                        |
            |  RELATÓRIO DE                          |  bloco creme
            |  ANÁLISE                               |  (60% altura)
            |  PROCESSUAL                            |
            |                                        |
            |  EDIÇÃO EXECUTIVA                      |
            |                                        |
            |  ─── AUTOS                             |
            |      0001234-56.2024...                |
            |  ─── CLASSE                            |
            |      Cumprimento de sentença           |
            |  ─── VALOR                             |
            |      R$ 125.430,00                     |
            |                                        |
            +========== faixa ocre ==================+
            |                                        |
            |  TJ          TECJUSTICA                |
            |  monograma   Assessoria Judicial       |
            |              com IA                    |  bloco navy
            |                                        |  (40% altura)
            |              PREPARADO POR             |
            |              autor · data              |
            +----------------------------------------+
        """
        PAGE_W = 21.0
        CREME_H = 16.8
        FAIXA_H = 0.6
        NAVY_H = 11.2

        table = self.doc.add_table(rows=3, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(PAGE_W)
        _set_table_width(table, PAGE_W)
        _set_table_indent(table, 0)

        _set_row_height(table.rows[0], CREME_H, rule="exact")
        _set_row_height(table.rows[1], FAIXA_H, rule="exact")
        _set_row_height(table.rows[2], NAVY_H, rule="exact")

        # ============================================================
        # BLOCO 1: topo creme com título editorial
        # ============================================================
        top = table.cell(0, 0)
        top.width = Cm(PAGE_W)
        _set_cell_bg(top, HEX_CREAM)
        _set_cell_margins(top, top=800, bottom=400, left=900, right=900)
        _set_cell_borders(top)
        top.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        self._build_cover_top(top)

        # ============================================================
        # BLOCO 2: faixa horizontal ocre sólida (separador visual)
        # ============================================================
        faixa = table.cell(1, 0)
        faixa.width = Cm(PAGE_W)
        _set_cell_bg(faixa, HEX_OCRE)
        _set_cell_margins(faixa, top=0, bottom=0, left=0, right=0)
        _set_cell_borders(faixa)
        p = faixa.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(1)

        # ============================================================
        # BLOCO 3: base navy com monograma e institucional
        # ============================================================
        bottom = table.cell(2, 0)
        bottom.width = Cm(PAGE_W)
        _set_cell_bg(bottom, HEX_NAVY_DEEP)
        _set_cell_margins(bottom, top=500, bottom=500, left=900, right=900)
        _set_cell_borders(bottom)
        bottom.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self._build_cover_bottom(bottom)

    def _build_cover_top(self, cell: _Cell) -> None:
        """Monta o conteúdo do bloco creme superior da capa."""
        # 1. Eyebrow + document ID
        first = cell.paragraphs[0]
        first.paragraph_format.space_before = Pt(0)
        first.paragraph_format.space_after = Pt(0)
        eyebrow_text = self.eyebrow
        if self.numero_documento:
            eyebrow_text = f"{self.eyebrow}   ·   {self.numero_documento}"
        _styled_run(
            first, eyebrow_text, font=FONT_SANS, size=8,
            color=OCRE, bold=True, tracking=60, uppercase=True,
        )

        # 2. Espaço antes do título
        spacer = cell.add_paragraph()
        spacer.paragraph_format.space_before = Pt(18)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.add_run("").font.size = Pt(1)

        # 3. Título gigantesco em caps (56pt Caladea bold)
        title_lines = self._smart_split_title(self.titulo.upper())
        for line in title_lines:
            tp = cell.add_paragraph()
            tp.paragraph_format.space_after = Pt(0)
            tp.paragraph_format.line_spacing = 0.95
            _styled_run(
                tp, line, font=FONT_DISPLAY, size=52,
                color=NAVY_DEEP, bold=True, tracking=20,
            )

        # 4. Subtítulo estilo "TEMPLATE"
        sub_p = cell.add_paragraph()
        sub_p.paragraph_format.space_before = Pt(10)
        sub_p.paragraph_format.space_after = Pt(0)
        subtitle_upper = (self.subtitulo or "EDIÇÃO EXECUTIVA").upper()
        if len(subtitle_upper) > 40:
            subtitle_upper = "EDIÇÃO EXECUTIVA"
        _styled_run(
            sub_p, subtitle_upper, font=FONT_SANS, size=11,
            color=MUTED, bold=True, tracking=50,
        )

        # 5. Metadata em linhas horizontais (label + valor)
        if self.metadata:
            meta_spacer = cell.add_paragraph()
            meta_spacer.paragraph_format.space_before = Pt(24)
            meta_spacer.paragraph_format.space_after = Pt(0)
            meta_spacer.add_run("").font.size = Pt(1)
            self._cover_metadata_rows(cell, self.metadata[:4])

    def _cover_metadata_rows(self, parent_cell: _Cell, items: Sequence[tuple[str, str]]) -> None:
        """Linhas horizontais de metadata estilo ficha técnica."""
        for idx, (label, value) in enumerate(items):
            line_p = parent_cell.add_paragraph()
            line_p.paragraph_format.space_before = Pt(6) if idx > 0 else Pt(0)
            line_p.paragraph_format.space_after = Pt(0)
            self._paragraph_bottom_border(line_p, color=HEX_HAIRLINE, size=6)

            _styled_run(
                line_p, label.upper(), font=FONT_SANS, size=8,
                color=MUTED, bold=True, tracking=50,
            )
            line_p.add_run("    ")
            is_mono = any(ch.isdigit() for ch in value) and any(
                ch in value for ch in ".-/R$"
            )
            _styled_run(
                line_p, value,
                font=FONT_MONO if is_mono and len(value) <= 35 else FONT_DISPLAY,
                size=11 if is_mono else 13,
                color=NAVY_DEEP, bold=not is_mono,
            )

    def _build_cover_bottom(self, cell: _Cell) -> None:
        """Bloco navy inferior: monograma + texto institucional."""
        # Tabela 1x2: monograma à esquerda + textos à direita
        inner = cell.add_table(rows=1, cols=2)
        inner.autofit = False
        inner.columns[0].width = Cm(5.5)
        inner.columns[1].width = Cm(12.0)

        # --- Monograma "TJ" ---
        mono_cell = inner.cell(0, 0)
        mono_cell.width = Cm(5.5)
        _set_cell_margins(mono_cell, top=0, bottom=0, left=0, right=200)
        _set_cell_borders(mono_cell, right=(HEX_OCRE, 8))
        mono_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        mono_p = mono_cell.paragraphs[0]
        mono_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mono_p.paragraph_format.space_before = Pt(0)
        mono_p.paragraph_format.space_after = Pt(0)
        mono_p.paragraph_format.line_spacing = 1.0
        _styled_run(
            mono_p, "TJ", font=FONT_DISPLAY, size=110,
            color=OCRE_LIGHT, bold=True,
        )

        # --- Textos institucionais ---
        text_cell = inner.cell(0, 1)
        text_cell.width = Cm(12.0)
        _set_cell_margins(text_cell, top=0, bottom=0, left=400, right=0)
        _set_cell_borders(text_cell)
        text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        name_p = text_cell.paragraphs[0]
        name_p.paragraph_format.space_before = Pt(0)
        name_p.paragraph_format.space_after = Pt(0)
        _styled_run(
            name_p, "TECJUSTICA", font=FONT_SANS, size=18,
            color=WHITE, bold=True, tracking=40,
        )

        tag_p = text_cell.add_paragraph()
        tag_p.paragraph_format.space_before = Pt(2)
        tag_p.paragraph_format.space_after = Pt(0)
        _styled_run(
            tag_p, "Assessoria Judicial com Inteligência Artificial",
            font=FONT_BODY, size=11, color=OCRE_LIGHT, italic=True,
        )

        # Separador interno
        sep_p = text_cell.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(14)
        sep_p.paragraph_format.space_after = Pt(0)
        self._paragraph_bottom_border(sep_p, color=HEX_OCRE, size=6)
        sep_p.add_run("").font.size = Pt(1)

        label_p = text_cell.add_paragraph()
        label_p.paragraph_format.space_before = Pt(10)
        label_p.paragraph_format.space_after = Pt(0)
        _styled_run(
            label_p, "PREPARADO POR", font=FONT_SANS, size=7,
            color=OCRE_LIGHT, bold=True, tracking=60,
        )

        author_p = text_cell.add_paragraph()
        author_p.paragraph_format.space_before = Pt(2)
        author_p.paragraph_format.space_after = Pt(0)
        _styled_run(
            author_p, self.autor, font=FONT_DISPLAY, size=13,
            color=WHITE, bold=True,
        )

        class_p = text_cell.add_paragraph()
        class_p.paragraph_format.space_before = Pt(6)
        class_p.paragraph_format.space_after = Pt(0)
        _styled_run(
            class_p, self.classificacao, font=FONT_MONO, size=8,
            color=OCRE_LIGHT, bold=True, tracking=30,
        )

    @staticmethod
    def _smart_split_title(title: str) -> list[str]:
        """Quebra o título em 2-3 linhas buscando respiração semântica."""
        words = title.split()
        if len(words) <= 2:
            return [title]
        if len(words) == 3:
            return [words[0], " ".join(words[1:])]
        # Para títulos de 4+ palavras, tenta dividir próximo ao meio
        mid = len(words) // 2
        for offset in range(3):
            idx = mid + offset
            if 0 < idx < len(words):
                return [" ".join(words[:idx]), " ".join(words[idx:])]
        return [title]

    @staticmethod
    def _paragraph_bottom_border(paragraph, color: str, size: int = 6) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        existing = p_pr.find(qn("w:pBdr"))
        if existing is not None:
            p_pr.remove(existing)
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    # ---- TOC ----

    def add_toc(self) -> None:
        self.add_heading("Sumário", level=1, numeral="—")
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        _set_run_font(run, FONT_BODY)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY
        _insert_field(
            run,
            'TOC \\o "1-3" \\h \\z \\u',
            placeholder="Clique com o botão direito e escolha 'Atualizar campo' para carregar o sumário.",
        )
        self._page_break()

    # ---- Headings ----

    def add_heading(self, text: str, level: int = 1, numeral: str | None = None) -> None:
        if level == 1:
            self._heading_level_1(text, numeral=numeral or "")
        elif level == 2:
            self._heading_level_2(text)
        else:
            self._heading_level_3(text)

    def _heading_level_1(self, text: str, numeral: str = "") -> None:
        self._vertical_space(Pt(14))

        if numeral:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            _styled_run(p, numeral, font=FONT_MONO, size=10, color=OCRE, bold=True, tracking=30)

        p = self.doc.add_heading(level=1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        _set_run_font(run, FONT_DISPLAY)
        run.font.size = Pt(26)
        run.font.color.rgb = NAVY_DEEP
        run.bold = True

        self._hairline(width_cm=6, color=HEX_OCRE)
        self._vertical_space(Pt(4))

    def _heading_level_2(self, text: str) -> None:
        p = self.doc.add_heading(level=2)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        _set_run_font(run, FONT_DISPLAY)
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        run.bold = True
        run.italic = True

    def _heading_level_3(self, text: str) -> None:
        p = self.doc.add_heading(level=3)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        _styled_run(p, text, font=FONT_SANS, size=10, color=NAVY, bold=True, tracking=40, uppercase=True)

    # ---- Paragraphs ----

    def add_paragraph(self, text: str, *, justify: bool = True) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        _styled_run(p, text, font=FONT_BODY, size=11, color=BODY)

    def add_lead(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        _styled_run(p, text, font=FONT_DISPLAY, size=14, color=NAVY, italic=True)

    def add_bullets(self, items: Iterable[str]) -> None:
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            for r in p.runs:
                r.clear()
            run = p.add_run(item)
            _set_run_font(run, FONT_BODY)
            run.font.size = Pt(11)
            run.font.color.rgb = BODY

    # ---- Blocks ----

    def add_code(self, code: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        _set_cell_bg(cell, HEX_CODE_BG)
        _set_cell_margins(cell, top=120, bottom=120, left=200, right=200)
        _set_cell_borders(cell, left=(HEX_OCRE, 18))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for i, line in enumerate(code.strip("\n").splitlines()):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            _set_run_font(run, FONT_MONO)
            run.font.size = Pt(9)
            run.font.color.rgb = NAVY_DEEP
        self.doc.add_paragraph()

    def add_callout(self, text: str, kind: str = "info") -> None:
        palette = {
            "info": (HEX_NAVY, "NOTA", NAVY),
            "warn": (HEX_OCRE, "ATENÇÃO", OCRE),
            "ok": ("046A38", "CONFIRMADO", RGBColor(0x04, 0x6A, 0x38)),
        }
        bar_hex, label, color = palette.get(kind, palette["info"])
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        _set_cell_bg(cell, HEX_CREAM)
        _set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
        _set_cell_borders(cell, left=(bar_hex, 24), top=(HEX_HAIRLINE, 4), bottom=(HEX_HAIRLINE, 4))

        label_p = cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(2)
        _styled_run(label_p, label, font=FONT_SANS, size=8, color=color, bold=True, tracking=50, uppercase=True)

        body_p = cell.add_paragraph()
        body_p.paragraph_format.space_after = Pt(0)
        _styled_run(body_p, text, font=FONT_BODY, size=10, color=BODY)
        self.doc.add_paragraph()

    def add_quote(self, text: str, author: str = "") -> None:
        self._vertical_space(Pt(6))

        quote_p = self.doc.add_paragraph()
        quote_p.paragraph_format.left_indent = Cm(0.8)
        quote_p.paragraph_format.right_indent = Cm(0.8)
        quote_p.paragraph_format.space_after = Pt(4)
        _add_left_border(quote_p, HEX_OCRE, size=18)
        _styled_run(quote_p, f"\u201C{text}\u201D", font=FONT_DISPLAY, size=13, color=NAVY_DEEP, italic=True)

        if author:
            attr_p = self.doc.add_paragraph()
            attr_p.paragraph_format.left_indent = Cm(0.8)
            attr_p.paragraph_format.space_before = Pt(2)
            attr_p.paragraph_format.space_after = Pt(8)
            _styled_run(attr_p, "— ", font=FONT_SANS, size=8, color=MUTED, tracking=80)
            _styled_run(attr_p, author, font=FONT_SANS, size=8, color=MUTED, bold=True, tracking=40, uppercase=True)

        self._vertical_space(Pt(4))

    # ---- Data cards ----

    def add_data_cards(self, items: Sequence[tuple[str, str]], columns: int = 2) -> None:
        rows = (len(items) + columns - 1) // columns
        table = self.doc.add_table(rows=rows, cols=columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_w = Cm(16.5 / columns)
        for col in table.columns:
            col.width = col_w

        for idx, (label, value) in enumerate(items):
            r, c = divmod(idx, columns)
            cell = table.rows[r].cells[c]
            cell.width = col_w
            _set_cell_bg(cell, HEX_CREAM)
            _set_cell_margins(cell, top=200, bottom=200, left=220, right=160)
            is_last_row = r == rows - 1
            is_last_col = c == columns - 1
            _set_cell_borders(
                cell,
                top=(HEX_HAIRLINE, 6),
                bottom=(HEX_OCRE, 8) if is_last_row else (HEX_HAIRLINE, 4),
                left=None,
                right=None if is_last_col else (HEX_HAIRLINE, 4),
            )

            label_p = cell.paragraphs[0]
            label_p.paragraph_format.space_after = Pt(4)
            _styled_run(label_p, label, font=FONT_SANS, size=7, color=MUTED, bold=True, tracking=50, uppercase=True)

            value_p = cell.add_paragraph()
            value_p.paragraph_format.space_after = Pt(0)
            is_mono = any(ch.isdigit() for ch in value) and any(ch in value for ch in "./-R$")
            _styled_run(
                value_p, value,
                font=FONT_MONO if is_mono and len(value) <= 30 else FONT_DISPLAY,
                size=11 if is_mono else 14,
                color=NAVY_DEEP,
                bold=not is_mono,
            )

        self.doc.add_paragraph()

    # ---- Tables ----

    def add_table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[str]],
        *,
        col_widths_cm: Sequence[float] | None = None,
    ) -> None:
        cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        if col_widths_cm:
            widths = [Cm(w) for w in col_widths_cm]
        else:
            widths = [Cm(16.5 / cols)] * cols

        header_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            cell = header_cells[i]
            cell.width = widths[i]
            _set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
            _set_cell_borders(
                cell,
                top=(HEX_NAVY, 12),
                bottom=(HEX_NAVY, 6),
                left=None,
                right=None,
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _styled_run(p, text, font=FONT_SANS, size=8, color=NAVY_DEEP, bold=True, tracking=50, uppercase=True)

        for r_idx, row in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            is_last = r_idx == len(rows) - 1
            for c_idx, value in enumerate(row):
                cell = row_cells[c_idx]
                cell.width = widths[c_idx]
                _set_cell_margins(cell, top=110, bottom=110, left=120, right=120)
                _set_cell_borders(
                    cell,
                    top=None,
                    bottom=(HEX_NAVY, 8) if is_last else (HEX_HAIRLINE, 4),
                    left=None,
                    right=None,
                )
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                _styled_run(p, value, font=FONT_BODY, size=10, color=BODY)

        self.doc.add_paragraph()

    # ---- Timeline ----

    def add_timeline(self, events: Sequence[tuple[str, str]]) -> None:
        table = self.doc.add_table(rows=len(events), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        date_w = Cm(3.2)
        bullet_w = Cm(0.9)
        text_w = Cm(12.4)

        for row_idx, (date, desc) in enumerate(events):
            is_last = row_idx == len(events) - 1

            date_cell = table.rows[row_idx].cells[0]
            date_cell.width = date_w
            _set_cell_margins(date_cell, top=80, bottom=80, left=0, right=120)
            _set_cell_borders(date_cell, bottom=(HEX_HAIRLINE, 4) if not is_last else None)
            p = date_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _styled_run(p, date, font=FONT_MONO, size=9, color=NAVY, bold=True, tracking=40)

            bullet_cell = table.rows[row_idx].cells[1]
            bullet_cell.width = bullet_w
            _set_cell_margins(bullet_cell, top=80, bottom=80, left=60, right=60)
            _set_cell_borders(bullet_cell, bottom=(HEX_HAIRLINE, 4) if not is_last else None)
            bp = bullet_cell.paragraphs[0]
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(0)
            _styled_run(bp, "◆", font=FONT_DISPLAY, size=10, color=OCRE, bold=True)

            text_cell = table.rows[row_idx].cells[2]
            text_cell.width = text_w
            _set_cell_margins(text_cell, top=80, bottom=80, left=120, right=0)
            _set_cell_borders(text_cell, bottom=(HEX_HAIRLINE, 4) if not is_last else None)
            tp = text_cell.paragraphs[0]
            tp.paragraph_format.space_before = Pt(0)
            tp.paragraph_format.space_after = Pt(0)
            _styled_run(tp, desc, font=FONT_BODY, size=10, color=BODY)

        self.doc.add_paragraph()

    # ---- KPI ----

    def add_kpi(self, items: Sequence[tuple[str, str, str]]) -> None:
        """Cards de KPI grandes: label em caps + numeral display + descritor."""
        cols = len(items)
        table = self.doc.add_table(rows=1, cols=cols)
        table.autofit = False
        col_w = Cm(16.5 / cols)
        for col in table.columns:
            col.width = col_w

        for idx, (label, value, desc) in enumerate(items):
            cell = table.rows[0].cells[idx]
            cell.width = col_w
            _set_cell_margins(cell, top=200, bottom=200, left=160, right=160)
            _set_cell_borders(
                cell,
                top=(HEX_OCRE, 12),
                bottom=(HEX_HAIRLINE, 4),
                left=None if idx == 0 else (HEX_HAIRLINE, 4),
                right=None,
            )

            label_p = cell.paragraphs[0]
            label_p.paragraph_format.space_after = Pt(4)
            _styled_run(label_p, label, font=FONT_SANS, size=7, color=MUTED, bold=True, tracking=50, uppercase=True)

            value_p = cell.add_paragraph()
            value_p.paragraph_format.space_after = Pt(2)
            _styled_run(value_p, value, font=FONT_DISPLAY, size=26, color=NAVY_DEEP, bold=True)

            desc_p = cell.add_paragraph()
            desc_p.paragraph_format.space_after = Pt(0)
            _styled_run(desc_p, desc, font=FONT_BODY, size=9, color=MUTED, italic=True)

        self.doc.add_paragraph()

    # ---- Signature ----

    def add_signature(self, name: str, role: str = "", local_data: str = "") -> None:
        self._vertical_space(Pt(20))

        if local_data:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            _styled_run(p, local_data, font=FONT_BODY, size=11, color=BODY, italic=True)

        line = self.doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(2)
        _styled_run(line, "_" * 50, font=FONT_BODY, size=10, color=BODY)

        n = self.doc.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        n.paragraph_format.space_after = Pt(0)
        _styled_run(n, name, font=FONT_DISPLAY, size=13, color=NAVY_DEEP, bold=True)

        if role:
            r_p = self.doc.add_paragraph()
            r_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_p.paragraph_format.space_before = Pt(0)
            _styled_run(r_p, role, font=FONT_SANS, size=8, color=MUTED, bold=True, tracking=50, uppercase=True)

    # ---- Layout utilities ----

    def _hairline(self, width_cm: float = 16.5, color: str = HEX_OCRE_LIGHT) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(width_cm)
        cell = table.cell(0, 0)
        cell.width = Cm(width_cm)
        _set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        _set_cell_borders(cell, top=(color, 6))
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(1)

    def _vertical_space(self, size: Pt) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = size
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(1)

    def _page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- Export ----

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)
        return path

    def export_pdf(self, docx_path: str | Path, pdf_path: str | Path | None = None) -> Path:
        docx_path = Path(docx_path)
        out_dir = Path(pdf_path).parent if pdf_path else docx_path.parent
        out_dir.mkdir(parents=True, exist_ok=True)
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                str(docx_path),
                "--outdir",
                str(out_dir),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice falhou ao converter {docx_path}:\n{result.stderr}"
            )
        default_pdf = out_dir / f"{docx_path.stem}.pdf"
        if pdf_path:
            target = Path(pdf_path)
            if default_pdf != target:
                default_pdf.rename(target)
            return target
        return default_pdf
